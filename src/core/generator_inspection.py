"""
generator_inspection.py — 報價單 → 驗機單 轉換 v7

Excel 部分：
  1. 清除 K 欄（col>=11）
  2. 標題「報價單」→「驗機單」
  3. 清除指定欄位值
  4. 動態定位單價/小計欄並歸零（含公式格）
  5. A1 插入 Logo
  6. 「訂購請簽章回傳」下方加底線

Word 部分：
  - 依報價單「序」數量，每個品項各生一份 Word
  - 套用 template/template.docx
  - 品號插入「改造」前面
  - 該品項「品名/規格」欄中所有紅色文字插入改造行下方
"""

import re, shutil, zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

from openpyxl import load_workbook
from openpyxl.cell import MergedCell

from _paths import TEMPLATE_DIR
from pathlib import Path as _Path
OUTPUT_DIR = _Path(r"Z:\Mika\驗收單及改造記錄單\Quoteflow_output")
WORD_TEMPLATE = TEMPLATE_DIR / "template.docx"
TITLE_RE      = re.compile(r'報[\s　]*價[\s　]*單')
CLEAR_KW      = {"電話", "傳真", "聯絡人", "統一編號", "聯絡地址", "EMAIL"}
ZERO_KW       = {"營業稅", "應收總金額", "合計"}

_NS = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'


# ── 共用工具 ─────────────────────────────────────────────────────────

def _norm(s) -> str:
    return re.sub(r'[\s　\-：:]+', '', str(s or '')).upper()

def _strip_ws(s) -> str:
    return re.sub(r'[\s　]+', '', str(s or ''))

def _cell_str(cell) -> str:
    if cell is None or isinstance(cell, MergedCell):
        return ''
    v = cell.value
    return '' if v is None else str(v).strip()

def _zero(v):
    if isinstance(v, (int, float)):
        return 0
    if isinstance(v, str) and re.sub(r'[NT$,\s]', '', v).replace('.', '', 1).isdigit():
        return "NT$0"
    return v



# ── 紅色文字抽取（by row） ───────────────────────────────────────────

def _is_red_color(color_el) -> bool:
    if color_el is None:
        return False
    rgb = color_el.get('rgb', '').upper()
    if len(rgb) == 8:
        r, g, b = int(rgb[2:4], 16), int(rgb[4:6], 16), int(rgb[6:8], 16)
    elif len(rgb) == 6:
        r, g, b = int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16)
    else:
        return False
    return r > 180 and g < 80 and b < 80


def _extract_red_text_by_row(src_path: str) -> dict:
    """
    解析 xlsx XML，回傳 {row_number: [紅色文字行, ...]}。
    掃描所有欄位，以 ※ 前綴過濾，避免因欄位位置不同而漏抓。
    """
    tag = lambda n: f'{{{_NS}}}{n}'

    shared        = []   # list of [(text, is_red), ...]
    red_style_idxs = set()   # cellXfs 索引中字色為紅的

    with zipfile.ZipFile(src_path, 'r') as zf:
        names = zf.namelist()

        # ── 1. styles.xml：找出哪些 cellXfs 索引的字色是紅的 ──────
        if 'xl/styles.xml' in names:
            with zf.open('xl/styles.xml') as f:
                st_root = ET.parse(f).getroot()
            # 建立 fontId → is_red 對應表
            font_is_red = {}
            fonts_el = st_root.find(tag('fonts'))
            if fonts_el is not None:
                for i, font_el in enumerate(fonts_el.findall(tag('font'))):
                    color = font_el.find(tag('color'))
                    font_is_red[i] = _is_red_color(color)
            # 掃描 cellXfs，收集紅色 style index
            cell_xfs = st_root.find(tag('cellXfs'))
            if cell_xfs is not None:
                for idx, xf in enumerate(cell_xfs.findall(tag('xf'))):
                    fid = xf.get('fontId')
                    if fid is not None and font_is_red.get(int(fid), False):
                        red_style_idxs.add(idx)

        # ── 2. sharedStrings.xml：解析 rich text runs ────────────
        if 'xl/sharedStrings.xml' in names:
            with zf.open('xl/sharedStrings.xml') as f:
                root = ET.parse(f).getroot()
            for si in root.findall(tag('si')):
                runs = []
                for r_el in si.findall(tag('r')):
                    rpr   = r_el.find(tag('rPr'))
                    t_el  = r_el.find(tag('t'))
                    txt   = (t_el.text or '') if t_el is not None else ''
                    color = rpr.find(tag('color')) if rpr is not None else None
                    runs.append((txt, _is_red_color(color)))
                if not runs:
                    t_el = si.find(tag('t'))
                    if t_el is not None:
                        runs.append((t_el.text or '', False))
                shared.append(runs)

        # ── 3. sheet XML ──────────────────────────────────────────
        sheet_path = next(
            (n for n in names if re.match(r'xl/worksheets/sheet\d+\.xml', n)),
            None
        )
        if not sheet_path:
            return {}

        with zf.open(sheet_path) as f:
            ws_root = ET.parse(f).getroot()

    result = {}
    for row_el in ws_root.iter(tag('row')):
        row_num = int(row_el.get('r', 0))
        for c_el in row_el.findall(tag('c')):
            # 儲存格層級是否為紅色 style
            cell_red = int(c_el.get('s', -1)) in red_style_idxs

            t_attr = c_el.get('t', '')
            runs   = []

            if t_attr == 's':
                v_el = c_el.find(tag('v'))
                if v_el is not None:
                    idx = int(v_el.text)
                    if idx < len(shared):
                        runs = shared[idx]
            elif t_attr == 'inlineStr':
                is_el = c_el.find(tag('is'))
                if is_el is not None:
                    for r_el2 in is_el.findall(tag('r')):
                        rpr   = r_el2.find(tag('rPr'))
                        t_el  = r_el2.find(tag('t'))
                        txt   = (t_el.text or '') if t_el is not None else ''
                        color = rpr.find(tag('color')) if rpr is not None else None
                        runs.append((txt, _is_red_color(color)))

            red_lines = []
            for text, run_red in runs:
                if run_red or cell_red:   # run 層級 OR 儲存格層級皆算紅
                    for line in text.splitlines():
                        line = line.strip()
                        if line and line.startswith('※'):
                            red_lines.append(line)
            if red_lines:
                result[row_num] = result.get(row_num, []) + red_lines

    return result


# ── Word 生成（套模板）──────────────────────────────────────────────

_W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
_A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'


def _fill_textbox(doc, first_char: str):
    """
    在右下內框（面積較小的 wps:wsp）填入 first_char，
    字型自動貼合方框大小，水平+垂直置中。
    """
    from lxml import etree

    wsps = doc.element.body.findall(f'.//{{{_WPS}}}wsp')
    if not wsps or not first_char:
        return

    def _area(wsp):
        spPr = wsp.find(f'{{{_WPS}}}spPr')
        if spPr is None:
            return 0
        xfrm = spPr.find(f'{{{_A}}}xfrm')
        if xfrm is None:
            return 0
        ext = xfrm.find(f'{{{_A}}}ext')
        if ext is None:
            return 0
        return int(ext.get('cx', 0)) * int(ext.get('cy', 0))

    inner = max(wsps, key=_area)   # 面積最大 = 外框（要填字的那個）

    # 垂直置中
    body_pr = inner.find(f'{{{_WPS}}}bodyPr')
    if body_pr is not None:
        body_pr.set('anchor', 'ctr')

    txbx = inner.find(f'{{{_WPS}}}txbx')
    if txbx is None:
        return
    tc = txbx.find(f'{{{_W}}}txbxContent')
    if tc is None:
        tc = etree.SubElement(txbx, f'{{{_W}}}txbxContent')

    # 清除舊內容
    for child in list(tc):
        tc.remove(child)

    SZ_HALF_PT = '96'    # 48pt × 2

    p   = etree.SubElement(tc,  f'{{{_W}}}p')
    pPr = etree.SubElement(p,   f'{{{_W}}}pPr')
    jc  = etree.SubElement(pPr, f'{{{_W}}}jc')
    jc.set(f'{{{_W}}}val', 'center')
    spc = etree.SubElement(pPr, f'{{{_W}}}spacing')
    spc.set(f'{{{_W}}}before', '0')
    spc.set(f'{{{_W}}}after',  '0')

    r    = etree.SubElement(p,   f'{{{_W}}}r')
    rPr  = etree.SubElement(r,   f'{{{_W}}}rPr')
    sz   = etree.SubElement(rPr, f'{{{_W}}}sz')
    sz.set(f'{{{_W}}}val', SZ_HALF_PT)
    szCs = etree.SubElement(rPr, f'{{{_W}}}szCs')
    szCs.set(f'{{{_W}}}val', SZ_HALF_PT)

    t = etree.SubElement(r, f'{{{_W}}}t')
    t.text = first_char


def _generate_word_from_template(part_no: str, red_lines: list,
                                 out_path: Path, customer: str = '') -> Path:
    """
    複製 template.docx，做三件事：
    1. 右下內框填入客戶名稱第一字（72pt，置中貼合）
    2. 品號插入「改造」前面
    3. 改造行正下方插入「長___寬___高___」（2pt）；再下方插入紅色※文字
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shutil.copy(str(WORD_TEMPLATE), str(out_path))
    doc = Document(str(out_path))

    # ── ① 填入客戶第一字 ─────────────────────────────────────────
    _fill_textbox(doc, customer[0] if customer else '')

    # ── ② 找「改造」段落 ──────────────────────────────────────────
    kaizao_para = None
    for para in doc.paragraphs:
        if '改造' in para.text:
            kaizao_para = para
            break

    if kaizao_para is None:
        doc.save(str(out_path))
        return out_path

    # 品號插到「改造」前
    for run in kaizao_para.runs:
        if '改造' in run.text:
            run.text = f'{part_no} ' + run.text
            break

    ref_p = kaizao_para._p

    # ── ③ 插入紅色 ※ 文字（逆序 addnext → 正序排列，固定 28pt）──
    for line in reversed(red_lines):
        new_p    = OxmlElement('w:p')
        new_r    = OxmlElement('w:r')
        new_rPr  = OxmlElement('w:rPr')
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), 'C0392B')
        new_rPr.append(color_el)
        sz_el = OxmlElement('w:sz')
        sz_el.set(qn('w:val'), '52')        # 52 half-pt = 26pt
        new_rPr.append(sz_el)
        new_r.append(new_rPr)
        new_t = OxmlElement('w:t')
        new_t.text = line.replace('※', '□', 1)
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        ref_p.addnext(new_p)

    # ── ④ 插入「長___寬___高___」緊接在改造後（最後 addnext = 最前）
    dim_p   = OxmlElement('w:p')
    dim_r   = OxmlElement('w:r')
    dim_rPr = OxmlElement('w:rPr')
    dim_sz  = OxmlElement('w:sz')
    dim_sz.set(qn('w:val'), '56')   # 56 half-pt = 28pt
    dim_szCs = OxmlElement('w:szCs')
    dim_szCs.set(qn('w:val'), '56')
    dim_rPr.append(dim_sz)
    dim_rPr.append(dim_szCs)
    dim_r.append(dim_rPr)
    dim_t = OxmlElement('w:t')
    dim_t.text = '長______寬______高______'
    dim_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    dim_r.append(dim_t)
    dim_p.append(dim_r)
    ref_p.addnext(dim_p)

    # ── ④-b 插入紅色「附配件 □電線/□充電器」在長寬高之後 ─────────
    acc_p    = OxmlElement('w:p')
    acc_r    = OxmlElement('w:r')
    acc_rPr  = OxmlElement('w:rPr')
    acc_clr  = OxmlElement('w:color')
    acc_clr.set(qn('w:val'), 'C0392B')
    acc_rPr.append(acc_clr)
    acc_sz   = OxmlElement('w:sz')
    acc_sz.set(qn('w:val'), '52')   # 26pt
    acc_rPr.append(acc_sz)
    acc_r.append(acc_rPr)
    acc_t    = OxmlElement('w:t')
    acc_t.text = '附配件 □電線/□充電器'
    acc_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    acc_r.append(acc_t)
    acc_p.append(acc_r)
    dim_p.addnext(acc_p)

    # ── ⑤ 調整緩衝空行，讓「插入項目 + size-28 空行」維持 8 行 ──────
    # 緩衝區共 10 行：前 8 行 size-28（可調整）、後 2 行保留不動
    # 插入了 1 行長寬高 + 1 行附配件 + N 行 ※，需從前 8 行移除 N+2 行
    n_to_remove = min(len(red_lines) + 2, 8)

    in_zone        = False
    buffer_empties = []
    for para in doc.paragraphs:
        if kaizao_para is not None and para._p is kaizao_para._p:
            in_zone = True
            continue
        if '主管簽核' in para.text:
            break
        if in_zone and not para.text.strip():
            buffer_empties.append(para)

    # 只移除前 8 行（size-28 群），後 2 行（buffer_empties[8:]）保留
    for para in buffer_empties[:n_to_remove]:
        para._p.getparent().remove(para._p)

    doc.save(str(out_path))
    return out_path


# ── 主函式 ───────────────────────────────────────────────────────────

def generate_inspection(src_path: str, data: dict, output_dir: _Path | None = None):
    """
    回傳 (excel_path, [word_path, ...]) tuple。
    """
    _out = output_dir or OUTPUT_DIR
    _out.mkdir(parents=True, exist_ok=True)

    customer = data["header"].get("customer", "客戶")
    out_path = _out / f"驗機單-{customer}.xlsx"
    shutil.copy(src_path, out_path)

    wb = load_workbook(out_path)
    ws = wb.active

    # ── ① 清除 K 欄（col >= 11）──────────────────────────────────
    from openpyxl.styles import PatternFill, Border, Font, Alignment
    _no_fill   = PatternFill(fill_type=None)
    _no_border = Border()
    _def_font  = Font()
    _def_align = Alignment()

    for row in ws.iter_rows():
        for cell in row:
            if cell.column >= 11 and not isinstance(cell, MergedCell):
                cell.value     = None
                cell.fill      = _no_fill
                cell.border    = _no_border
                cell.font      = _def_font
                cell.alignment = _def_align

    def _img_col(img) -> int:
        a = img.anchor
        if isinstance(a, str):
            m = re.match(r'([A-Za-z]+)', a)
            if m:
                col = 0
                for ch in m.group(1).upper():
                    col = col * 26 + (ord(ch) - 64)
                return col
        else:
            try:
                return img.anchor._from.col + 1
            except Exception:
                pass
        return 0

    ws._images = [img for img in ws._images if _img_col(img) < 11]

    # ── ② 標題替換 ────────────────────────────────────────────────
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell, MergedCell) or not cell.value:
                continue
            if _strip_ws(str(cell.value)) == "報價單":
                cell.value = TITLE_RE.sub('驗 機 單', str(cell.value))
            elif "報價單號" in str(cell.value):
                cell.value = str(cell.value).replace("報價單號", "驗機單號")

    # ── ③ 清除指定欄位值 ──────────────────────────────────────────
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell, MergedCell) or not cell.value:
                continue
            if any(kw in _norm(str(cell.value)) for kw in CLEAR_KW):
                for c in range(cell.column + 1, cell.column + 8):
                    t = ws.cell(row=cell.row, column=c)
                    if not isinstance(t, MergedCell) and t.value:
                        t.value = None
                        break

    # ── ④ 動態定位單價/小計/品名欄，歸零 ────────────────────────
    def _is_numeric(v):
        if isinstance(v, (int, float)):
            return True
        if isinstance(v, str):
            s = re.sub(r'[NT$,\s]', '', v).replace('.', '', 1).lstrip('-')
            return bool(s) and s.isdigit()
        return False

    def _should_clear(v):
        if v is None:
            return False
        if isinstance(v, (int, float)):
            return True
        if isinstance(v, str):
            return v.startswith('=') or _is_numeric(v)
        return False

    price_col    = 8
    subtotal_col = 9
    spec_col     = 3
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell, MergedCell) or not cell.value:
                continue
            text = _norm(str(cell.value))
            if text == "單價" and price_col == 8:
                price_col = cell.column
            if text in ("小計", "金額") and subtotal_col == 9:
                subtotal_col = cell.column
            if "品名" in text and spec_col == 3:
                spec_col = cell.column
        if price_col != 8 or subtotal_col != 9:
            break

    for row in ws.iter_rows():
        rn = row[0].row

        is_total_row = any(
            not isinstance(cell, MergedCell) and cell.value and
            any(kw in _norm(str(cell.value)) for kw in ZERO_KW)
            for cell in row
        )

        if is_total_row:
            for c in range(2, 12):
                t = ws.cell(row=rn, column=c)
                if isinstance(t, MergedCell) or t.value is None:
                    continue
                new_v = _zero(t.value)
                if new_v != t.value:
                    t.value = new_v
        else:
            cp = ws.cell(row=rn, column=price_col)
            if not isinstance(cp, MergedCell) and _should_clear(cp.value):
                cp.value = None

            cs = ws.cell(row=rn, column=subtotal_col)
            if not isinstance(cs, MergedCell) and _should_clear(cs.value):
                cs.value = 0

    # ── ⑤ 掃描品項行（seq, part_no, row_num）────────────────────
    item_header_row = -1
    for row in ws.iter_rows():
        for cell in row:
            if not isinstance(cell, MergedCell) and _norm(str(cell.value or '')) in ('序', '序號'):
                item_header_row = cell.row
                break
        if item_header_row > 0:
            break

    item_rows = []   # [(seq, part_no, row_num)]
    if item_header_row > 0:
        for r in range(item_header_row + 1, item_header_row + 300):
            a_val = ws.cell(row=r, column=1).value
            if a_val is None:
                continue
            a_str = str(a_val).strip()
            if '應收總金額' in a_str or ('合計' in a_str and r > item_header_row + 2):
                break
            if a_str.isdigit():
                part_no = _cell_str(ws.cell(row=r, column=2))
                item_rows.append((int(a_str), part_no, r))

    # ── ⑥ 「訂購請簽章回傳」下方加底線 ──────────────────────────
    from openpyxl.styles import Side
    sig_row = -1
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell, MergedCell) or not cell.value:
                continue
            if "訂購" in str(cell.value) and "簽章" in str(cell.value):
                sig_row = cell.row
                sig_col = cell.column
                break
        if sig_row > 0:
            break

    if sig_row > 0:
        sig_cell = ws.cell(row=sig_row + 4, column=sig_col)
        if not isinstance(sig_cell, MergedCell):
            sig_cell.border = Border(bottom=Side(style='thin'))

    # ── ⑦ 掃描 A 欄，找到「備註:」→ 清空該格、右邊同列、右邊往下7列 ──
    def _clear(ws, row, col):
        cell = ws.cell(row=row, column=col)
        if not isinstance(cell, MergedCell):
            cell.value = None
            return
        for merge in ws.merged_cells.ranges:
            if merge.min_row <= row <= merge.max_row and merge.min_col <= col <= merge.max_col:
                ws.cell(row=merge.min_row, column=merge.min_col).value = None
                return

    for row in ws.iter_rows(min_col=1, max_col=1):
        cell = row[0]
        if isinstance(cell, MergedCell) or cell.value is None:
            continue
        if '備註' in str(cell.value):
            r, c = cell.row, cell.column
            cell.value = None
            for dr in range(0, 8):
                _clear(ws, r + dr, c + 1)
            for dr in range(4, 8):
                _clear(ws, r + dr, c + 2)

    wb.save(out_path)

    # ── ⑧ 產生 Word（每個品項一份）──────────────────────────────
    word_paths = []
    if WORD_TEMPLATE.exists() and item_rows:
        try:
            red_by_row = _extract_red_text_by_row(src_path)

            for i, (seq, part_no, row_num) in enumerate(item_rows):
                end_row = item_rows[i + 1][2] if i + 1 < len(item_rows) else row_num + 200
                red_lines = []
                for r in range(row_num, end_row):
                    red_lines.extend(red_by_row.get(r, []))

                safe_part = re.sub(r'[\\/:*?"<>|]', '_', part_no) if part_no else f'item{seq}'
                date_tag  = datetime.today().strftime("%Y%m%d")
                word_out  = _out / f"{customer}{safe_part}改造紀錄單-{date_tag}.docx"
                _generate_word_from_template(part_no, red_lines, word_out,
                                             customer=customer)
                word_paths.append(word_out)
        except Exception as e:
            print(f"[警告] Word 檔生成失敗：{e}")

    return out_path, word_paths


if __name__ == "__main__":
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    from parser import parse

    if len(sys.argv) < 2:
        print("用法：python generator_inspection.py 報價單.xlsx")
        sys.exit(1)

    src = sys.argv[1]
    if not Path(src).exists():
        print(f"找不到檔案：{src}")
        sys.exit(1)

    data = parse(src)
    xlsx, words = generate_inspection(src, data)
    print(f"驗機單 Excel：{xlsx}")
    for w in words:
        print(f"驗機單 Word ：{w}")
