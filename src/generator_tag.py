"""
generator_tag.py — 生成維修掛件 (template_fix.docx)
"""
import shutil
from datetime import datetime
from pathlib import Path

from _paths import TEMPLATE_DIR, OUTPUT_DIR

TAG_TEMPLATE = TEMPLATE_DIR / "template_fix.docx"

_W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
_A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'


def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _set_para_text(para, new_text):
    """將段落第一個 run 設為 new_text，其餘 run 清空（保留第一個 run 的樣式）。"""
    if not para.runs:
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _fill_keyword(para, keyword, value, prepend=False):
    """
    找到 keyword 後：
    - prepend=False：將 keyword 之後的文字替換為 value
    - prepend=True ：將 value 插到 keyword 前面
    """
    text = _para_text(para)
    if keyword not in text:
        return False
    if prepend:
        new_text = text.replace(keyword, value + keyword, 1)
    else:
        idx = text.index(keyword) + len(keyword)
        new_text = text[:idx] + value
    _set_para_text(para, new_text)
    return True


def _iter_all_paras(doc):
    """yield 文件中所有段落：表格格子 → 一般段落 → 文字框內段落。"""
    from docx.text.paragraph import Paragraph as _Para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    yield from doc.paragraphs
    # 文字框（wps:txbx）內的段落不在 doc.paragraphs 裡，需手動萃取
    for txbx in doc.element.body.iter(f'{{{_WPS}}}txbx'):
        tc = txbx.find(f'{{{_W}}}txbxContent')
        if tc is not None:
            for p_elem in tc.findall(f'{{{_W}}}p'):
                yield _Para(p_elem, doc)


def _fill_customer_char(doc, char: str):
    """在右下文字框（面積最大的 wps:wsp）填入 char，size 36。"""
    if not char:
        return
    from lxml import etree

    wsps = doc.element.body.findall(f'.//{{{_WPS}}}wsp')
    if not wsps:
        return

    def _area(wsp):
        spPr = wsp.find(f'{{{_WPS}}}spPr')
        if spPr is None:
            return 0
        xfrm = spPr.find(f'{{{_A}}}xfrm')
        if xfrm is None:
            return 0
        ext = xfrm.find(f'{{{_A}}}ext')
        if ext is None:
            return 0
        return int(ext.get('cx', 0)) * int(ext.get('cy', 0))

    target = max(wsps, key=_area)

    body_pr = target.find(f'{{{_WPS}}}bodyPr')
    if body_pr is not None:
        body_pr.set('anchor', 'ctr')

    txbx = target.find(f'{{{_WPS}}}txbx')
    if txbx is None:
        return
    tc = txbx.find(f'{{{_W}}}txbxContent')
    if tc is None:
        tc = etree.SubElement(txbx, f'{{{_W}}}txbxContent')

    for child in list(tc):
        tc.remove(child)

    SZ = '72'  # 36pt × 2
    p    = etree.SubElement(tc,  f'{{{_W}}}p')
    pPr  = etree.SubElement(p,   f'{{{_W}}}pPr')
    jc   = etree.SubElement(pPr, f'{{{_W}}}jc')
    jc.set(f'{{{_W}}}val', 'center')
    spc  = etree.SubElement(pPr, f'{{{_W}}}spacing')
    spc.set(f'{{{_W}}}before', '0')
    spc.set(f'{{{_W}}}after',  '0')
    r    = etree.SubElement(p,   f'{{{_W}}}r')
    rPr  = etree.SubElement(r,   f'{{{_W}}}rPr')
    sz   = etree.SubElement(rPr, f'{{{_W}}}sz')
    sz.set(f'{{{_W}}}val', SZ)
    szCs = etree.SubElement(rPr, f'{{{_W}}}szCs')
    szCs.set(f'{{{_W}}}val', SZ)
    t    = etree.SubElement(r,   f'{{{_W}}}t')
    t.text = char


def generate_tag(data, tag_data, output_dir=None):
    """
    data     : 報價單解析資料（取客戶名稱）
    tag_data : {
        no             : str  — No. 後面的數字
        part_no        : str  — 品號（放在「(序:)」前面）
        seq_no         : str  — 序號
        problem        : str  — 問題描述
        pullback_date  : str  — 拉回日期（YYYY/MM/DD）
        repair_status  : str  — 維修狀況
    }
    """
    from docx import Document

    out_dir = Path(output_dir) if output_dir else OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    customer   = data["header"].get("customer", "客戶")
    date_tag   = datetime.today().strftime("%Y%m%d")
    out_path   = out_dir / f"維修掛件-{customer}-{date_tag}.docx"

    shutil.copy(str(TAG_TEMPLATE), str(out_path))
    doc = Document(str(out_path))

    no_val        = str(tag_data.get("no", ""))
    part_no       = str(tag_data.get("part_no", ""))
    seq_no        = str(tag_data.get("seq_no", ""))
    problem       = str(tag_data.get("problem", ""))
    pullback_date = str(tag_data.get("pullback_date", ""))
    repair_status = str(tag_data.get("repair_status", ""))

    all_paras      = list(_iter_all_paras(doc))
    para_originals = [_para_text(p) for p in all_paras]

    _W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _max_font_size(para) -> int:
        """掃段落 XML 內所有 w:sz，回傳最大 pt 值（0 = 找不到）。"""
        sizes = []
        for sz in para._p.iter(f'{{{_W_NS}}}sz'):
            val = sz.get(f'{{{_W_NS}}}val')
            if val and val.isdigit():
                sizes.append(int(val) // 2)   # half-pt → pt
        return max(sizes, default=0)

    def _find_and_fill(keywords, value, prepend=False, max_pt=28):
        if not value:
            return
        for para, original in zip(all_paras, para_originals):
            sz = _max_font_size(para)
            if sz > max_pt:   # 0 = 未設定（繼承樣式），不過濾
                continue
            for kw in keywords:
                if kw in original:
                    _fill_keyword(para, kw, value, prepend=prepend)
                    return

    # ── No. ──────────────────────────────────────────────────────
    _find_and_fill(["No."], no_val, max_pt=999)

    # ── 品號 + 序號 → 同一段落 "(序:seq_no)" 前面加品號 ───────────
    for para, original in zip(all_paras, para_originals):
        for paren_kw in ["(序:)", "(序：)"]:
            if paren_kw in original:
                colon = paren_kw[2]
                new_text = original.replace(
                    paren_kw,
                    f"{part_no}(序{colon}{seq_no})" if seq_no else f"{part_no}{paren_kw}",
                    1
                )
                _set_para_text(para, new_text)
                break
        else:
            continue
        break

    # ── 其他欄位（含空格變體；裝飾性大字不含冒號所以不需過濾字體大小）
    _find_and_fill(["問題:", "問題：", "問題 :", "問題 ："],     problem,       max_pt=999)
    _find_and_fill(["維修狀況:", "維修狀況：", "維修狀況 :", "維修狀況 ："], repair_status, max_pt=999)

    # ── 拉回日期：固定用 size 24 的新 run 附加在「拉回 ：」後面 ───────
    if pullback_date:
        from lxml import etree
        _PULLBACK_KWS = ["拉回:", "拉回：", "拉回 :", "拉回 ："]
        for para, original in zip(all_paras, para_originals):
            if any(kw in original for kw in _PULLBACK_KWS):
                r    = etree.SubElement(para._p, f'{{{_W}}}r')
                rPr  = etree.SubElement(r,       f'{{{_W}}}rPr')
                sz   = etree.SubElement(rPr,     f'{{{_W}}}sz')
                sz.set(f'{{{_W}}}val', '48')       # 24pt × 2
                szCs = etree.SubElement(rPr,     f'{{{_W}}}szCs')
                szCs.set(f'{{{_W}}}val', '48')
                t    = etree.SubElement(r,       f'{{{_W}}}t')
                t.text = pullback_date
                break

    try:
        _fill_customer_char(doc, customer[0] if customer else "")
    except Exception:
        pass

    doc.save(str(out_path))
    return out_path
